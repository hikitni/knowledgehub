from __future__ import annotations

import json
import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from .models import ParsedDocument


class ParseError(RuntimeError):
    pass


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def _title_from_markdown(text: str, fallback: str) -> str:
    for line in text.splitlines():
        match = re.match(r"^#\s+(.+?)\s*$", line)
        if match:
            return match.group(1).strip()
    return fallback


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".markdown", ".txt", ".yaml", ".yml", ".json"}:
            text = _read_text(path)
            if suffix == ".json":
                # Validate JSON but preserve original formatting and line numbers.
                json.loads(text)
            title = _title_from_markdown(text, path.stem) if suffix in {".md", ".markdown"} else path.stem
            return ParsedDocument(title=title, content=text, file_type=suffix.lstrip("."))
        if suffix in {".html", ".htm"}:
            raw = _read_text(path)
            soup = BeautifulSoup(raw, "html.parser")
            title = soup.title.get_text(" ", strip=True) if soup.title else path.stem
            return ParsedDocument(title=title, content=soup.get_text("\n", strip=True), file_type="html")
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return ParsedDocument(title=path.stem, content="\n\n".join(pages), file_type="pdf")
        if suffix == ".docx":
            document = Document(str(path))
            paragraphs = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text for cell in row.cells))
            return ParsedDocument(title=path.stem, content="\n".join(paragraphs), file_type="docx")
    except Exception as exc:
        raise ParseError(f"Failed to parse {path.name}: {type(exc).__name__}") from exc
    raise ParseError(f"Unsupported file type: {suffix or '<none>'}")
